"""Build cp5_to_cp65_understanding.docx — run with plain python3."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(__file__), "cp5_to_cp65_understanding.docx")

doc = Document()


# ── style helpers ─────────────────────────────────────────────────────────────

def _set_font(run, name="Arial", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(text, level=1, size=None):
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    _set_font(run, size=size or sizes[level], bold=True)
    return p


def body(text, italic=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _set_font(run, size=10, italic=italic)
    return p


def code(text):
    """Monospaced code block."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    _set_font(run, size=10, bold=True)
    return p


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Arial"
    # Data rows
    for ri, row in enumerate(rows):
        drow = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = "Arial"
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def hr():
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(180, 180, 180)


# ── Title page ────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CP5 → CP6 → CP6.5\nNavigation Codebase — Complete Understanding Guide")
_set_font(run, size=20, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run(
    "NeuroGait Quadruped Navigation  |  Implementation branch\n"
    "Go2 robot · Isaac Lab · skrl PPO · CNN+MLP policy"
)
_set_font(run2, size=11, italic=True, color=(80, 80, 80))

doc.add_paragraph()
hr()
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — THE BIG PICTURE
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 1 — The Big Picture", level=1)

body(
    "The entire stack is a hierarchical control system: a high-level navigation "
    "policy (what direction should I go?) drives a frozen locomotion policy "
    "(how do I move my legs?). The two policies never see each other's internals — "
    "they communicate only through a 3D velocity command: [vx, vy, heading]."
)

code("""\
                    TRAINING LOOP  (scripts/train.py)
                           │
                 gym.make("NeuroGait-Navigation-CP65-v0")
                           │
              ┌────────────┼────────────┐
              ▼            ▼            ▼
         SceneCfg    ObservationsCfg  RewardsCfg
         (obstacles,  (what policy    (what behavior
          cameras,     sees)           is rewarded)
          robot)           │                │
              │            ▼                ▼
              │     NavigationPolicy   9 reward terms
              │     (CNN+MLP, 1615→3)  (multiplicative core)
              │            │
              │            ▼
              │     PreTrainedPolicyAction
              │     (frozen 12-joint loco, 235→12)
              │            │
              └────────────┼────────────┘
                           ▼
                    Isaac Sim / PhysX
                    (50 Hz physics,
                     5 Hz nav steps via decimation=40)
""")

body(
    "IMPORTANT: The navigation policy runs at 5 Hz (every 40 physics steps = 0.2 s). "
    "The locomotion policy runs at 50 Hz internally. Isaac Lab's 'decimation' parameter "
    "controls this: the manager-based env calls env.step() every 40 physics steps."
)

doc.add_paragraph()


# ── Section 1b: Two-level hierarchy ──────────────────────────────────────────

heading("The Two-Level Hierarchy", level=2)

code("""\
Level 1 — Navigation Policy (TRAINED HERE, 5 Hz)
  Input:  1615-dim obs  [40×40 grid | 9 waypoints | 3 velocity | 3 gravity]
  Output: 3-dim action  [vx ∈ [-1,1] m/s | vy ∈ [-1,1] m/s | heading ∈ [-π,π] rad]
  Model:  CNN branch (grid → 64-dim) + MLP branch (scalars → 32-dim) → 3

Level 2 — Locomotion Policy (FROZEN, 50 Hz, pre-trained Go2 checkpoint)
  Input:  235-dim obs   [joint positions, velocities, body velocity, commands, ...]
  Output: 12-dim        [target joint angles for 12 Go2 joints]
  Loaded: PreTrainedPolicyActionCfg in navigation_env_cfg.py (path to .pt file)
""")

body(
    "The navigation policy's [vx, vy, heading] output is injected into the "
    "locomotion policy's 'velocity_commands' observation slot by "
    "PreTrainedPolicyAction. The locomotion policy THINKS a joystick is giving "
    "it commands — it has no idea a navigation policy is in the loop."
)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — EVERY FILE EXPLAINED
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 2 — Every File Explained", level=1)


# ── config/go2/ ──────────────────────────────────────────────────────────────

heading("config/go2/ — Environment Configuration", level=2)

body(
    "These files define WHAT the training environment looks like. They are pure "
    "Python dataclasses decorated with @configclass — no computation happens here. "
    "Isaac Lab reads them at gym.make() time and builds the actual env."
)

label("navigation_env_cfg.py")
body("The main config file. Contains the full inheritance chain from the base env "
     "down to CP65. Every @configclass here is one registered task.")
code("""\
ManagerBasedRLEnvCfg                        (Isaac Lab)
  └── NavigationBaseEnvCfg                  (navigation_base_env_cfg.py)
        └── NeuroGaitNavigationGo2BaseEnvCfg
              ├── NeuroGaitNavigationCP1EnvCfg     (CP3/CP4 rule-based)
              └── NeuroGaitNavigationCP5EnvCfg     (CP5 — trained policy)
                    ├── NeuroGaitNavigationCP5EnvCfg_PLAY
                    └── NeuroGaitNavigationCP6EnvCfg     (CP6 — A* replan)
                          ├── NeuroGaitNavigationCP6EnvCfg_PLAY
                          └── NeuroGaitNavigationCP65EnvCfg  (CP6.5 — generator)
                                └── NeuroGaitNavigationCP65EnvCfg_PLAY
""")
body("Key things set in __post_init__:")
body("• RayCasterCamera (depth sensor) pointed downward at 30° pitch", indent=1)
body("• PreTrainedPolicyActionCfg — loads frozen locomotion checkpoint", indent=1)
body("• ObservationsCfg — wires occupancy_grid_obs_cp5 + future_waypoints_obs + velocity + gravity", indent=1)
body("• RewardsCfg — either CP5RewardsCfg or CP6RewardsCfg", indent=1)
body("• EventsCfg — reset event that either randomizes obstacles (CP6) or generates scene (CP6.5)", indent=1)

doc.add_paragraph()

label("cp5_rewards.py")
body("Dataclass holding reward term configs for CP5. Separated from the env config "
     "to keep navigation_env_cfg.py clean. Contains 7 terms (additive).")
table(
    ["Term", "Function", "Weight"],
    [
        ["velocity_toward_goal", "cp5_reward_velocity_toward_goal", "+10.0"],
        ["goal_proximity", "cp5_reward_goal_proximity", "+0.1"],
        ["goal_reached", "cp5_reward_goal_reached", "+20.0"],
        ["collision", "cp5_penalty_collision_velocity_scaled", "-5.0"],
        ["stuck", "cp5_penalty_stuck", "-0.3"],
        ["heading", "cp5_reward_heading", "+0.1"],
        ["smoothness", "cp5_penalty_smoothness", "-0.01"],
    ],
    [2.2, 2.8, 0.8],
)

label("cp6_rewards.py")
body("CP6 reward terms (9 terms, multiplicative core). Replaces additive CP5 "
     "velocity reward with a product of forward/lateral/heading Gaussians — "
     "prevents crab-walking.")
table(
    ["Term", "Function", "Weight"],
    [
        ["navigation_core", "cp6_reward_navigation_core", "+10.0"],
        ["path_following", "cp6_reward_path_following", "+0.5"],
        ["goal_proximity", "cp5_reward_goal_proximity", "+0.1"],
        ["goal_reached", "cp5_reward_goal_reached", "+50.0"],
        ["slow_near_goal", "cp6_reward_slow_near_goal", "+3.0"],
        ["graduated_clearance", "cp6_penalty_graduated_clearance", "-0.05"],
        ["collision", "cp5_penalty_collision_velocity_scaled", "-1.5"],
        ["stuck", "cp6_penalty_stuck_v2", "-0.3"],
        ["smoothness", "cp6_penalty_smoothness_2nd_order", "-1.0"],
    ],
    [2.2, 2.8, 0.8],
)

label("__init__.py  (config/go2/)")
body("Registers all tasks with gymnasium. The key line is gym.register(id=..., "
     "env_cfg_entry_point=..., skrl_cfg_entry_point=...). Task IDs registered:")
code("""\
NeuroGait-Navigation-Unitree-Go2-Play-v1    (CP1 play only)
NeuroGait-Navigation-CP5-v0                 (CP5 train)
NeuroGait-Navigation-CP5-Play-v0            (CP5 play)
NeuroGait-Navigation-CP6-v0                 (CP6 train)
NeuroGait-Navigation-CP6-Play-v0            (CP6 play)
NeuroGait-Navigation-CP65-v0                (CP6.5 train)  ← train this
NeuroGait-Navigation-CP65-Play-v0           (CP6.5 play)
""")

doc.add_paragraph()


# ── mdp/ ─────────────────────────────────────────────────────────────────────

heading("mdp/ — Observations, Rewards, Events", level=2)

body(
    "These files contain the actual COMPUTATION. Functions here are referenced "
    "by the config dataclasses in config/go2/ via ObsTerm, RewTerm, EventTerm, "
    "and DoneTerm wrappers. Isaac Lab calls them every step."
)

label("mdp/observations.py")
body("All observation functions. The most important ones for CP5/CP6/CP6.5:")

code("""\
occupancy_grid_obs_cp5(env) → Tensor (E, 1600)
  Converts depth camera output to a 40×40 binary occupancy grid.
  Robot is at cell (row=10, col=20) — asymmetric: 2m behind, 6m ahead.
  Depth values outside [0.05, 2.0]m are masked.
  Cells are 0.20m × 0.20m → total coverage: 8m × 8m.

future_waypoints_obs(env) → Tensor (E, 9)
  Returns the NEXT 3 waypoints relative to robot body frame.
  Each waypoint = (dx_body, dy_body, dist) → 3 scalars × 3 waypoints = 9.
  Advances waypoint index when robot is within 0.3m of current waypoint.

robot_velocity_obs(env) → Tensor (E, 3)
  Returns [vx, vy, yaw_rate] in body frame from robot state.

_cp5_init_waypoint_state(env) → None
  Lazy init: called by EVERY reward and obs function that needs waypoints.
  On first call: runs A* (CP5/CP6) or detects generator has already run (CP6.5).
  Stores env._cp5_waypoints (E, W, 2) world coords as env attribute.
  If env._cp5_waypoints already exists (generator set it) — returns immediately.
  This is the KEY integration point between generator and rewards.

_cp5_reset_waypoint_state(env, env_ids) → None
  Called by EventTerm on episode reset to clear wp_idx, prev_dist, action history.
""")

label("mdp/rewards.py")
body("All CP5 and CP6 reward functions. Shared helper: _cp5_current_wp_and_dist() "
     "reads env._cp5_waypoints and returns current waypoint + distance.")

code("""\
CP5 reward functions:
  cp5_reward_velocity_toward_goal  → cos(Δθ) × vx × (1 + 1/(1+2d²))
  cp5_reward_goal_proximity        → (1-tanh(d/5)) + (1-tanh(d/1))  [dual-scale]
  cp5_reward_goal_reached          → 1.0 if dist < 0.3 m else 0.0
  cp5_penalty_collision_velocity_scaled → (1+4‖v‖²) × contact_bool
  cp5_penalty_stuck                → 1.0 if max_disp < 0.1m AND forward_cmd
  cp5_reward_heading               → cos(heading_error)
  cp5_penalty_smoothness           → ‖action_t - action_{t-1}‖

CP6 reward functions:
  cp6_reward_navigation_core       → exp(-Δfwd²/σ) × exp(-vy²/σ) × exp(-Δθ²/σ)
  cp6_reward_path_following        → exp(-dist_to_path² / 1.0)
  cp6_penalty_graduated_clearance  → sum over obstacles of exp(-d_i² / 0.25)
  cp6_reward_slow_near_goal        → (1 - speed.clamp(1)) × 𝟙[d < 1.5m]
  cp6_penalty_stuck_v2             → stuck AND NOT near_goal
  cp6_penalty_smoothness_2nd_order → ‖action_t - 2×action_{t-1} + action_{t-2}‖
""")

label("mdp/events.py")
body("Event functions called on episode reset. Two key ones for CP6/CP6.5:")

code("""\
cp6_randomize_obstacles_and_replan(env, env_ids, position_range)
  CP6 only. On reset: shifts ALL 9 obstacles by a shared random (dx, dy),
  runs A* to replan path, broadcasts waypoints to all envs. One-time log.

cp65_reset_with_generated_scene(env, env_ids, ramp_iterations=40_000)
  CP6.5 only. On reset:
  1. Calls env._curriculum.step() to advance difficulty
  2. Calls generate_scene() → (obstacles, waypoints, goal_xy)
  3. Calls apply_scene_to_env() to move rigid objects in sim
  4. Stores waypoints as env._cp5_waypoints (E, W, 2) world tensor
  5. Resets wp_idx, prev_dist, action history for resetting envs
  6. Rate-limited log (1 per 60s) showing curriculum progress + goal
""")

label("mdp/terminations.py")
body("Termination functions. The one that matters for navigation:")
code("""\
cp6_goal_reached(env, threshold=0.1) → BoolTensor (E,)
  True when robot is within 0.1m of env._cp5_waypoints[:, -1, :].
  The last waypoint is ALWAYS the goal (A* ends there, generator ends there).
  Works for both fixed goal (CP6) and random goal (CP6.5) without changes.
""")

label("mdp/__init__.py")
body("Exports everything. Key import chain: wildcard from isaaclab.envs.mdp "
     "(base locomotion functions) plus explicit imports of all navigation-specific "
     "functions. This lets navigation_env_cfg.py do: import mdp as nav_mdp.")

doc.add_paragraph()


# ── models/ ──────────────────────────────────────────────────────────────────

heading("models/ — Neural Network Architecture", level=2)

label("models/navigation_policy.py")
body("Defines NavigationPolicy (actor) and NavigationValue (critic) for skrl PPO.")

code("""\
NavigationPolicy  (inherits GaussianMixin, Model)
  │
  ├── CNN branch — processes 40×40 occupancy grid
  │     Conv2d(1→16,  5×5, stride=2, pad=2)   40×40 → 20×20  [ELU]
  │     Conv2d(16→32, 3×3, stride=2, pad=1)   20×20 → 10×10  [ELU]
  │     Conv2d(32→32, 3×3, stride=1, pad=1)   10×10 → 10×10  [ELU]
  │     Flatten → 3200
  │     Linear(3200→128) [ELU]
  │     Linear(128→64)   [ELU]
  │     Output: (B, 64)
  │
  ├── MLP branch — processes 15 scalar observations
  │     Linear(15→64) [ELU]
  │     Linear(64→32) [ELU]
  │     Output: (B, 32)
  │
  ├── Fusion head
  │     cat(64, 32) → (B, 96)
  │     Linear(96→128) [ELU]
  │     Linear(128→64) [ELU]
  │     Linear(64→3)
  │     Output: (B, 3)  ← raw logits
  │
  └── Output squashing
        vx, vy  → tanh(raw[:, :2])        ∈ [-1, 1] m/s
        heading → π × tanh(raw[:, 2:3])   ∈ [-π, π] rad

NavigationValue  (same CNN+MLP, Linear(64→1) instead of Linear(64→3))

Key skrl API note:
  compute(inputs, role) reads inputs["observations"]  (NOT inputs["states"])
  Returns: (mean_actions, {"log_std": self.log_std_parameter})
  GaussianMixin uses this dict key — returning 3-tuple causes KeyError.
""")

doc.add_paragraph()


# ── planning/ ────────────────────────────────────────────────────────────────

heading("planning/ — A* Path Planning (CP5 / CP6)", level=2)

body("Used in CP5 and CP6. CP6.5 has its own A* embedded in scene_generator.py "
     "so the planning/ module is NOT called during CP6.5 training.")

label("planning/global_grid.py")
body("Builds a 2D binary occupancy grid from env.scene rigid object positions. "
     "Robot-centred: origin = robot_pos_w - half_grid_size.")
code("""\
build_global_grid(env, grid_resolution=0.2, grid_size=200)
  → (grid, origin, obstacle_info)

Grid convention:
  row ↔ Y axis    (row 0 = min Y = origin_y)
  col ↔ X axis    (col 0 = min X = origin_x)
  Each obstacle inflated by 0.30m (C-space mapping)
  Robot half-width 0.155m + safety 0.145m = 0.30m total

BUG FIXED: origin was hardcoded to (0,0) — robot at (35,-35) was
completely outside the grid. Fixed to: origin = robot_pos - half_grid.
""")

label("planning/astar.py")
body("8-connected A* implementation on a 2D numpy grid. "
     "Diagonal cost = sqrt(2) for proper Euclidean distance.")

label("planning/planner.py")
body("Wraps global_grid.py + astar.py into a single AStarPlanner class.")
code("""\
AStarPlanner.plan(start_world_xy, goal_world_xy) → list of (x, y)
  Returns waypoints in world coords, downsampled to every 5th cell.
  All print spam removed (was flooding terminal during training).
""")

doc.add_paragraph()


# ── scene/ ───────────────────────────────────────────────────────────────────

heading("scene/ — Scene Generation and Curriculum (CP6.5)", level=2)

body("New in CP6.5. Replaces A* replanning with obstacles-first generation. "
     "These files have ZERO Isaac Lab imports — they run standalone.")

label("scene/scene_generator.py")
body("Core generator. Obstacles placed randomly, BFS verifies connectivity, "
     "A* finds optimal path — all in local env coordinates.")

code("""\
GO2_WIDTH   = 0.31 m   (actual Go2 body width)
SAFETY_MARGIN = 0.10 m (per-side clearance)
MIN_GAP     = 0.51 m   (0.31 + 2×0.10 — just squeezable)

random_goal(start_xy, dist_range, angle_range) → (gx, gy)
  Random goal at variable distance (5-10m) and angle (±46°).

generate_scene(start_xy, goal_xy, num_obstacles, min_gap_width, ...)
  → (obstacles, waypoints, goal_xy)
  Algorithm:
    1. Build empty occupancy grid (0.1m resolution)
    2. For each obstacle attempt:
       a. Random position, size in arena
       b. Reject if too close to start/goal (< 1.2m)
       c. Reject if overlaps existing obstacle
       d. Stamp obstacle + C-space inflation onto test_grid
       e. BFS: if start↔goal still connected → accept, update grid
    3. Run A* on final grid → path_cells
    4. Downsample to ~1m waypoint spacing
    Returns: obstacles list, waypoints list, goal_xy tuple

apply_scene_to_env(env, obstacles, device) → None
  Moves Isaac Lab RigidObjects to generated positions.
  Broadcasts local coords to world: world_pos = env_origin + local_pos
  Parks unused objects at y=1000m.
  torch imported lazily (only needed inside sim).

_stamp(), _bfs_connected(), _astar() — internal helpers
""")

label("scene/curriculum.py")
body("5-axis linear curriculum driven by call count. "
     "Training: step() every reset. Viz: set current_iteration directly.")

code("""\
NavigationCurriculum(ramp_iterations=40_000)
  .step()                → increment current_iteration
  .get_difficulty()      → dict with 5 keys:
    min_gap_width: 2.00m → 0.51m   (corridor tightness)
    num_obstacles: 3     → 12      (scene density)
    arena_padding: 3.0m  → 1.5m   (how close obstacles are to path)
    goal_dist:    (6,7)  → (5,10)  (goal distance spread in m)
    goal_angle:   ±0.3   → ±0.8   (goal angle spread in rad)
  .progress_str()        → human-readable one-liner for logs

Timing math:
  512 envs × 120s episodes / 0.2s nav_dt = 600 steps/episode/env
  Events per training iter ≈ 512 × 24 / 600 ≈ 20
  ramp_iterations=40_000 ≈ 40_000/20 = 2000 training iterations
""")

doc.add_paragraph()


# ── navigation_base_env_cfg.py ───────────────────────────────────────────────

heading("navigation_base_env_cfg.py — Base Scene", level=2)

body("Defines the base scene: terrain (flat plane), robot (Go2 URDF), "
     "9 obstacle rigid bodies (cube_01..06, cyl_01..03), contact force sensor, "
     "RayCaster camera, height scanner. All configs live here as @configclass.")

doc.add_paragraph()


# ── control/, managers/ ──────────────────────────────────────────────────────

heading("control/ and managers/ — Legacy CP4 Files", level=2)

body("control/waypoint_controller.py: PD-based pure-pursuit controller from CP4. "
     "Not used in CP5+ training but kept as a CP4 baseline for comparison.")
body("managers/: Legacy manager-based wrappers from early CP development. "
     "Not used in CP5+ — replaced by mdp/ functions.")

doc.add_paragraph()


# ── scripts/ ─────────────────────────────────────────────────────────────────

heading("scripts/ — Training and Evaluation", level=2)

label("scripts/train.py")
body("Universal training script. Works with any registered task. "
     "Key features: auto-detects log directory from task name, "
     "bypasses skrl Runner (directly builds PPO + SequentialTrainer + "
     "NavigationPolicy + NavigationValue), loads pre-trained checkpoint if available.")
code("""\
python3 scripts/train.py --task NeuroGait-Navigation-CP65-v0 --num_envs 512 --headless
Log directory: logs/skrl/neurogait_navigation_cp65_v0/YYYY-MM-DD_HH-MM-SS/
""")

label("scripts/test_reward_scales.py")
body("Loads a task config, runs random actions for 200 steps, prints mean/std "
     "of every reward term. Use BEFORE training to check reward magnitudes are balanced.")
code("""\
python3 scripts/test_reward_scales.py --task NeuroGait-Navigation-CP65-v0 --num_envs 64 --headless
""")

label("scripts/cp5/play.py")
body("Runs a trained checkpoint in simulation with rendering. "
     "Loads NavigationPolicy from .pt file, runs inference loop.")

label("scripts/cp6/eval_metrics.py")
body("50-episode evaluator: computes success rate, path efficiency "
     "(A* path length / actual path length), collision count, steps-to-goal. "
     "Outputs JSON summary and optional comparison table with CP5 baseline.")

label("scripts/cp6/visualize_generated_scenes.py")
body("Standalone visualization: 6-panel matplotlib figure showing curriculum "
     "progression (0% → 100%). No Isaac Sim needed. "
     "Saves to maps/cp6_5_obstacles_first_difficulty.png.")
code("""\
python3 scripts/cp6/visualize_generated_scenes.py
""")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE EVOLUTION
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 3 — The Evolution: CP5 → CP6 → CP6.5", level=1)

table(
    ["What", "CP5", "CP6", "CP6.5"],
    [
        ["Task ID", "CP5-v0", "CP6-v0", "CP65-v0 ← train this"],
        ["Scene", "Fixed 9 obstacles", "Same 9, shifted ±1.5m", "Random obstacles-first generator"],
        ["Goal", "Fixed local (8, 0)", "Fixed local (8, 0)", "Random: 5-10m, ±46°"],
        ["A* timing", "Runs once at init", "Replans each reset batch", "Embedded in generator"],
        ["Traversability", "Not guaranteed", "Not guaranteed (fallback)", "BFS-verified by construction"],
        ["Rewards", "7 terms, additive", "9 terms, multiplicative core", "Same as CP6 (unchanged)"],
        ["Curriculum", "None", "None", "5-axis ramp over 40k resets"],
        ["Why it fails", "Crab-walk, no stop", "A* spam, fixed goal", "N/A (current)"],
    ],
    [1.5, 1.5, 1.5, 2.5],
)

heading("Why obstacles-first is better than path-first", level=2)
code("""\
Path-first (CP6.5 v1 — discarded):
  1. Generate curvy cubic-spline path start→goal
  2. Place obstacles OUTSIDE the protected corridor
  Problem: obstacles are DECORATION — they never block the path
           path curves for no reason (A* would take straighter route)
           training scenes don't look like deployment scenes

Obstacles-first (CP6.5 v2 — current):
  1. Place obstacles randomly in tight arena around start→goal
  2. BFS flood-fill after each: reject if path blocked
  3. A* on final grid → SAME algorithm as deployment
  Result: obstacles CREATE gaps the policy must thread
          A* path is tight and realistic
          training = deployment
""")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — DATA FLOW: ONE TRAINING STEP
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 4 — Data Flow: One Complete env.step()", level=1)

code("""\
RL trainer (skrl PPO) calls env.step(action)
  action shape: (512, 3) = [vx, vy, heading] per env

Inside env.step():

  ┌── ActionManager ──────────────────────────────────────────────────────┐
  │  PreTrainedPolicyAction receives action (512, 3)                       │
  │  Stores as _raw_actions — rewires velocity_commands obs slot           │
  │  Builds internal 235-dim locomotion obs (joint pos, vel, body vel ...) │
  │  Runs frozen loco policy: 235 → 12 joint targets                      │
  │  PD controllers: targets → torques → applied to Go2                   │
  │                                                                         │
  │  Inner loop × 40 physics steps (50 Hz physics × 0.02s = 1 step):      │
  │    PhysX integrates: F=ma → new joint pos, vel, body state            │
  └────────────────────────────────────────────────────────────────────────┘

After 40 physics steps = 1 navigation step (0.2s):

  ┌── ObservationManager ─────────────────────────────────────────────────┐
  │  occupancy_grid_obs_cp5():                                              │
  │    depth camera → threshold [0.05, 2.0m] → project to 40×40 grid     │
  │    robot at (row=10, col=20), 0.2m cells, 8m×8m coverage              │
  │    → (512, 1600) float32 tensor                                         │
  │                                                                         │
  │  future_waypoints_obs():                                                │
  │    read env._cp5_waypoints[env_id, wp_idx:wp_idx+3]                   │
  │    transform to body frame, compute distances                          │
  │    advance wp_idx if robot within 0.3m of current waypoint            │
  │    → (512, 9) float32 tensor                                            │
  │                                                                         │
  │  robot_velocity_obs() → (512, 3)                                        │
  │  gravity_obs()        → (512, 3)                                        │
  │  concatenate → (512, 1615) obs tensor                                   │
  └────────────────────────────────────────────────────────────────────────┘

  ┌── RewardManager ──────────────────────────────────────────────────────┐
  │  navigation_core:  r_fwd × r_lat × r_hdg (multiplicative)             │
  │  path_following:   exp(-dist_to_path² / 1.0)                           │
  │  goal_proximity:   dual-scale tanh to final waypoint                   │
  │  goal_reached:     1.0 if dist < 0.3m                                  │
  │  slow_near_goal:   (1-speed) × 𝟙[d < 1.5m]                            │
  │  graduated_clearance: -exp per close obstacle                          │
  │  collision:        -(1+4‖v‖²) × contact_bool                          │
  │  stuck:            -𝟙[no_movement AND fwd_command AND not_near_goal]  │
  │  smoothness:       -‖Δaction + Δaction_prev‖  (2nd order jerk)         │
  │  total_reward = Σ weight_i × term_i  → (512,) float32                 │
  └────────────────────────────────────────────────────────────────────────┘

  ┌── TerminationManager ─────────────────────────────────────────────────┐
  │  cp6_goal_reached:   dist_to_final_waypoint < 0.1m → done             │
  │  time_out:           step_count > max_episode_length → done           │
  └────────────────────────────────────────────────────────────────────────┘

Return to trainer: (obs, reward, terminated, truncated, info)
PPO updates policy every 24 × 512 = 12,288 nav steps
""")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DATA FLOW: EPISODE RESET (CP6.5)
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 5 — Data Flow: Episode Reset (CP6.5)", level=1)

code("""\
When any env terminates, Isaac Lab calls cp65_reset_with_generated_scene(env, env_ids)

Step 1: Curriculum
  env._curriculum.step()         → current_iteration++
  diff = curriculum.get_difficulty()
    → {min_gap_width, num_obstacles, arena_padding, goal_dist, goal_angle}

Step 2: Scene generation (pure Python, no GPU, ~1ms)
  obstacles, waypoints, goal_xy = generate_scene(
    start_xy=(0,0), goal_xy=None, **diff
  )
  Internally:
    a. Random goal via random_goal() — variable distance + angle
    b. Build empty grid (0.1m/cell) around tight arena
    c. Try placing num_obstacles random obstacles:
       - Reject if too close to start/goal
       - Reject if overlaps existing obstacle
       - BFS check: if start↔goal still reachable → accept
    d. A* on final grid → path_cells
    e. Downsample to ~1m waypoint spacing → waypoints list

Step 3: Apply to sim
  apply_scene_to_env(env, obstacles, device)
    For each of 9 scene rigid objects:
      if i < len(obstacles): world_pos = env_origin + local_pos
      else:                  world_pos[:, 1] = 1000m  (park it far away)
      write_root_pose_to_sim(pose, all_env_ids)

Step 4: Store waypoints
  local_wp = torch.tensor(waypoints)           # (W, 2) local
  env._cp5_waypoints = local_wp + env_origins  # (E, W, 2) world
  # Last entry _cp5_waypoints[:, -1, :] = goal world pos
  # → cp6_goal_reached termination works automatically

Step 5: Reset tracking state for env_ids
  env._cp5_wp_idx[env_ids]      = 0
  env._cp5_prev_dist[env_ids]   = inf
  env._cp5_prev_action[env_ids] = 0
  env._cp5_pos_history[env_ids] = 0

Step 6: Rate-limited log (max 1 per 60s)
  [CP6.5] Curriculum 35%: gap=1.51m, obs=6, pad=2.5m, angle=±0.5rad
          goal=(3.2, 4.8) | 9 waypoints | 6 obstacles
""")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — REWARD FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 6 — Reward Functions: Why Each Term", level=1)

body("All CP6 reward terms. CP5 is a subset (7 of these, with different weights).")

heading("cp6_reward_navigation_core  weight=+10.0", level=2)
code("r = exp(-Δfwd²/0.25) × exp(-3vy²) × exp(-Δθ²/0.25)")
body("Forward: Gaussian around target speed 0.8 m/s in the goal direction. "
     "Lateral: penalises sideways crab-walk (Go2 should face forward). "
     "Heading: penalises facing away from current waypoint. "
     "MULTIPLICATIVE: if ANY component is ~0, the whole reward is ~0. "
     "This is why CP5 crab-walked: additive rewards let the robot earn "
     "goal_proximity while drifting sideways.")
body("Source: Miki et al. 2022 (Science Robotics) — orthogonal velocity concept.", italic=True)
body("Too high: robot sprints toward waypoint ignoring obstacles. Too low: policy ignores direction.")

heading("cp6_reward_path_following  weight=+0.5", level=2)
code("r = exp(-min_dist_to_path² / 1.0)")
body("Gaussian proximity to the FULL A*/generator path, not just current waypoint. "
     "Teaches robot to thread THROUGH obstacle gaps instead of going around the "
     "entire obstacle field.")
body("Weight history: 5.0 → 1.0 → 0.5 (too high caused path-hugging, not goal-seeking).", italic=True)

heading("cp5_reward_goal_proximity  weight=+0.1", level=2)
code("r = (1 - tanh(d/5)) + (1 - tanh(d/1))")
body("Dual scale: 5m scale gives global gradient (robot always has incentive to get closer). "
     "1m scale gives fine-grained gradient near goal. "
     "Uses FINAL waypoint, not current — persistent gradient that never resets.")
body("Source: Li et al. 2025, Eq. 1.", italic=True)
body("Small weight (0.1) — just a shaping signal, not the primary reward.")

heading("cp5_reward_goal_reached  weight=+50.0", level=2)
code("r = 1.0 if dist < 0.3m else 0.0  (CP5: 0.3m, CP6: 0.1m threshold)")
body("Large sparse bonus at episode success. Without this, the policy has no "
     "incentive to actually REACH the goal vs circling near it. "
     "Weight 50.0 makes it worth ~5 seconds of max navigation_core reward.")
body("Source: X-Nav 2025 — one-time success signal.", italic=True)

heading("cp6_reward_slow_near_goal  weight=+3.0", level=2)
code("r = (1 - speed.clamp(1)) × 𝟙[dist < 1.5m]")
body("Near goal: reward is higher when robot is SLOW. "
     "Without this: robot oscillates past the goal at full speed, never triggering "
     "the 0.1m threshold. With this: robot decelerates to dock.")

heading("cp6_penalty_graduated_clearance  weight=-0.05", level=2)
code("r = Σᵢ exp(-dᵢ² / 0.25)  summed over all obstacles")
body("Light continuous penalty for being close to obstacles. "
     "Teaches wall-hugging avoidance without blocking progress. "
     "Weight history: -1.0 → -0.05. At -1.0 it was 1147× dominant "
     "because depth camera fires non-NaN values every step even in open space.")

heading("cp5_penalty_collision_velocity_scaled  weight=-1.5", level=2)
code("r = (1 + 4(vx² + vy² + ωz²)) × 𝟙[base_contact_force > 1N]")
body("Velocity-scaled: crashing at high speed costs more than gentle contact. "
     "sensor_cfg MUST specify body_names='base' (trunk only). "
     "If '.*' is used, all bodies are selected and the mask logic fails silently.")

heading("cp6_penalty_stuck_v2  weight=-0.3", level=2)
code("r = 𝟙[max_disp < 0.1m AND fwd_cmd AND NOT near_goal]")
body("Fires when robot commands forward motion but doesn't move. "
     "Near-goal exemption: near the goal, slow speed is desired (see slow_near_goal). "
     "Without near-goal patch: penalty fires right when the robot should be docking.")
body("Source: SEA-Nav 2026, Table III with near-goal patch.", italic=True)

heading("cp6_penalty_smoothness_2nd_order  weight=-1.0", level=2)
code("r = ‖action_t - 2×action_{t-1} + action_{t-2}‖  (jerk = 2nd derivative)")
body("Second-order penalty (vs CP5's first-order). "
     "Penalises change in velocity change, not just velocity change. "
     "Produces smoother, more physically plausible motions.")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — THE CNN+MLP NETWORK
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 7 — CNN+MLP Network: Layer by Layer", level=1)

body("Full forward pass with actual tensor shapes (batch size B = 512):")

code("""\
Input obs: (B, 1615)
  Split at index 1600:
    grid_flat  = obs[:, :1600]    → (B, 1600)
    scalars    = obs[:, 1600:]    → (B, 15)

━━━ CNN BRANCH ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  grid_2d = grid_flat.view(B, 1, 40, 40)          → (B, 1, 40, 40)

  Conv2d(1→16, k=5, stride=2, pad=2)  [ELU]
    B×1×40×40 → B×16×20×20

  Conv2d(16→32, k=3, stride=2, pad=1) [ELU]
    B×16×20×20 → B×32×10×10

  Conv2d(32→32, k=3, stride=1, pad=1) [ELU]
    B×32×10×10 → B×32×10×10

  Flatten
    B×32×10×10 → (B, 3200)

  Linear(3200→128) [ELU]             → (B, 128)
  Linear(128→64)   [ELU]             → (B, 64)   ← grid_feat

━━━ MLP BRANCH ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  scalars = obs[:, 1600:]
    [1600:1609] 3 future waypoints × 3 (dx_body, dy_body, dist)
    [1609:1612] robot velocity [vx, vy, yaw_rate] in body frame
    [1612:1615] projected gravity vector [gx, gy, gz]

  Linear(15→64) [ELU]                → (B, 64)
  Linear(64→32) [ELU]                → (B, 32)   ← scalar_feat

━━━ FUSION ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  merged = cat([grid_feat, scalar_feat], dim=-1)  → (B, 96)

  Linear(96→128)  [ELU]              → (B, 128)
  Linear(128→64)  [ELU]              → (B, 64)
  Linear(64→3)                       → (B, 3)    ← raw logits

━━━ OUTPUT SQUASHING ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  vx_vy   = tanh(raw[:, :2])        → [-1.0, 1.0] m/s
  heading = π × tanh(raw[:, 2:3])   → [-π, π] rad

  → (B, 3)  = [vx, vy, heading] per env

Total parameters: ~700K  (CNN: ~590K, FC: ~110K)
""")

body(
    "The CNN sees SPATIAL obstacle layout. The MLP sees WHERE TO GO (waypoints) "
    "and HOW FAST I'M MOVING (velocity). Fusing them lets the policy decide: "
    "'given this obstacle field AND these waypoints, what velocity should I command?'"
)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — BUGS AND LESSONS
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 8 — Bugs and Lessons Learned", level=1)

bugs = [
    (
        "1. heading_command vs yaw_rate mismatch (CP5 init)",
        "Robot spun in place.",
        "Navigation action[2] was interpreted as yaw_rate but locomotion policy "
        "expected a heading angle. Command interface mismatch.",
        "Changed action[2] to mean absolute heading (not yaw rate). "
        "Locomotion policy converts heading → yaw error internally."
    ),
    (
        "2. Camera GPU overflow (CP5 init)",
        "CUDA OOM on first reset.",
        "RayCasterCamera output was stored full-resolution; with 512 envs × "
        "200×200 depth image the buffer overflowed.",
        "Switched to 40×40 asymmetric grid (see concept/06_asymmetric_grid.md). "
        "Downsampling done on GPU before storing."
    ),
    (
        "3. Reward imbalance — smoothness 33× dominant (CP5 init)",
        "Policy refused to move — any action was immediately penalized.",
        "smoothness weight was -1.0 on a first-order penalty. "
        "Random actions have ‖Δa‖ ≈ 1.0 per step, so penalty was ≈ 1.0 "
        "while max reward was ≈ 0.03. Policy learned to stay still.",
        "Reduced smoothness to -0.01 (regularizer only). Run test_reward_scales.py "
        "BEFORE every training run to catch this class of bug."
    ),
    (
        "4. Progress always zero — prev_dist reset every step (CP5 init)",
        "velocity_toward_goal reward was always ~0.",
        "prev_dist was being reset to inf every step instead of persisting. "
        "The progress term (d_prev - d_curr) was always near zero.",
        "Fixed state management: prev_dist only resets on episode reset "
        "(in EventTerm), not on every step."
    ),
    (
        "5. render_interval < decimation (CP5 init)",
        "render_interval error from Isaac Lab on startup.",
        "render_interval must be a multiple of decimation (40). "
        "Had set render_interval=4 which is less than decimation.",
        "Set render_interval=40 (= one nav step). Documentation was wrong."
    ),
    (
        "6. skrl Runner expects class not instance (CP5 train)",
        "TypeError: Runner expects model class, got NavigationPolicy instance.",
        "PPORunnerCfg requires passing the CLASS (NavigationPolicy) not an "
        "instantiated object. Runner instantiates it internally with env spaces.",
        "Bypassed Runner entirely: use PPO + SequentialTrainer directly, "
        "pass model instances. Gives full control over model creation."
    ),
    (
        "7. inputs['states'] returns None in skrl (CP5 train)",
        "model.compute() received None for the observation tensor.",
        "skrl passes inputs['observations'] for on-policy PPO, not inputs['states']. "
        "Had written self.policy(inputs['states']) following wrong documentation.",
        "Changed to inputs['observations'] in NavigationPolicy.compute()."
    ),
    (
        "8. Model on CPU, input on GPU (CP5 play)",
        "RuntimeError: Expected all tensors on same device.",
        "play.py loaded checkpoint on CPU. env.step() returns obs on CUDA. "
        "model(obs.cpu()) worked but was slow and easy to miss.",
        "Added model.to(device) immediately after load_state_dict(). "
        "Check device placement whenever loading a checkpoint."
    ),
    (
        "9. policy.act() returns 2 values not 3 (CP5 play)",
        "ValueError: too many values to unpack.",
        "skrl GaussianMixin.act() returns (actions, log_prob) in eval mode, "
        "not (actions, log_prob, outputs) like training mode.",
        "Unpacked correctly: actions, _ = policy.act(inputs, role='policy')."
    ),
    (
        "10. A* grid at wrong location (CP6 BUG 1)",
        "A* reports 'no path found' for every env. Robot gets straight-line fallback.",
        "Grid origin was hardcoded to (0,0). With env_spacing=12m and 512 envs, "
        "robots are placed at e.g. (35, -35). The 8m×8m grid around (0,0) "
        "doesn't contain the robot at all.",
        "Fixed: origin = robot_pos_w[0, :2] - half_grid_size. "
        "The grid is now centred on the robot's actual world position."
    ),
    (
        "11. graduated_clearance 1147× dominant (CP6 BUG 2)",
        "Policy ran away from all obstacles and refused to approach them.",
        "Weight was -1.0. Depth camera fires non-NaN values every step even "
        "in open space (ambient depth values ~1.5m). "
        "With 9 obstacles each contributing exp(-d²/0.25) ≈ 0.127, total "
        "penalty was ~1.14 per step. navigation_core max ≈ 0.001.",
        "Reduced weight to -0.05 (22× reduction). Always test graduated_clearance "
        "with test_reward_scales.py before training."
    ),
    (
        "12. Robot crab-walks sideways (CP5 result)",
        "Policy translated laterally toward goal instead of facing it.",
        "CP5 used additive rewards. goal_proximity rewarded getting closer "
        "regardless of heading. vy was unconstrained so sideways was fine.",
        "CP6 multiplicative core: r_lateral = exp(-3vy²) collapses the "
        "entire navigation_core reward when vy ≠ 0."
    ),
    (
        "13. Robot takes long detours (CP5 result)",
        "Robot went around the entire obstacle field instead of through gaps.",
        "No reward for staying near the A* path. Robot maximized goal_proximity "
        "by taking the least-resistance route (around all obstacles).",
        "Added cp6_reward_path_following: rewards proximity to the A* path, "
        "which threads THROUGH the gaps."
    ),
    (
        "14. Robot doesn't stop at goal, oscillates (CP5 result)",
        "Robot reaches 0.5m from goal then oscillates back and forth.",
        "No deceleration incentive near goal. goal_reached sparse bonus "
        "requires < 0.3m but robot overshoots at full speed.",
        "Added cp6_reward_slow_near_goal (weight +3.0). Also tightened "
        "goal_reached threshold from 0.3m to 0.1m."
    ),
    (
        "15. Path-first generator: obstacles are decoration (CP6.5 v1)",
        "Visualization showed A* path was basically a straight line; "
        "obstacles were all far from the path.",
        "Generating path first then placing obstacles OUTSIDE the corridor "
        "means obstacles can never be near the path. "
        "The curriculum's corridor narrowing had no effect on difficulty.",
        "Switched to obstacles-first: place obstacles first, BFS checks "
        "connectivity, A* plans through the actual gaps."
    ),
    (
        "16. Fixed goal: policy memorizes direction (CP6.5 v1)",
        "Policy learned to always move in the +X direction regardless of "
        "where obstacles or waypoints were.",
        "Goal was always at local (8, 0) — always straight ahead. "
        "With enough training, policy learned 'go right' as a shortcut.",
        "Added random_goal(): variable distance (5-10m) and angle (±46°). "
        "Policy MUST read the waypoint obs to know where to go."
    ),
]

for bug_name, symptom, cause, fix in bugs:
    heading(bug_name, level=2)
    label("Symptom:")
    body(symptom, indent=0.5)
    label("Root cause:")
    body(cause, indent=0.5)
    label("Fix + lesson:")
    body(fix, indent=0.5)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — FILE DEPENDENCY GRAPH
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 9 — File Dependency Graph", level=1)

code("""\
CONFIGURATION LAYER
───────────────────
navigation_env_cfg.py
  ├── imports: isaaclab.managers (EventTerm, ObsTerm, RewTerm, DoneTerm)
  ├── imports: .cp5_rewards.CP5RewardsCfg
  ├── imports: .cp6_rewards.CP6RewardsCfg
  ├── imports: mdp as nav_mdp  (references functions by name in *Cfg objects)
  └── registered as gym tasks via config/go2/__init__.py

cp5_rewards.py / cp6_rewards.py
  └── imports: mdp as nav_mdp  (references reward functions)

config/go2/__init__.py
  └── imports: navigation_env_cfg (indirectly via gym string refs)

COMPUTATION LAYER (mdp/)
────────────────────────
mdp/__init__.py
  ├── wildcard from isaaclab.envs.mdp  (locomotion base functions)
  ├── from .observations import: occupancy_grid_obs_cp5, future_waypoints_obs,
  │                               robot_velocity_obs, _cp5_init_waypoint_state, ...
  ├── from .rewards import: all cp5_* and cp6_* reward functions
  ├── from .terminations import: cp6_goal_reached
  └── from .events import: cp6_randomize_obstacles_and_replan,
                           cp65_reset_with_generated_scene

mdp/observations.py
  ├── imports: planning.global_grid (for A* init)
  ├── imports: planning.planner (for A* init)
  └── _cp5_init_waypoint_state SKIPPED if env._cp5_waypoints already exists
      (generator sets it first → A* never called in CP6.5)

mdp/rewards.py
  ├── imports: observations._cp5_init_waypoint_state (called every step)
  └── reads env._cp5_waypoints, env._cp5_wp_idx, env._cp5_prev_dist (env attrs)

mdp/events.py
  ├── CP6:   imports planning.global_grid, planning.planner
  └── CP6.5: imports scene.scene_generator, scene.curriculum

SCENE GENERATION (scene/)
──────────────────────────
scene/scene_generator.py
  └── NO PROJECT IMPORTS (standalone — numpy + scipy only)
      torch imported lazily (inside apply_scene_to_env only)

scene/curriculum.py
  └── imports scene.scene_generator.MIN_GAP  (0.51m constant)

scene/__init__.py
  └── re-exports: generate_scene, apply_scene_to_env, NavigationCurriculum, ...

NEURAL NETWORK (models/)
────────────────────────
models/navigation_policy.py
  └── imports: torch, torch.nn, skrl.models.torch (Model, GaussianMixin, DeterministicMixin)
      NO project imports — can be imported standalone

PLANNING (planning/)
──────────────────────
planning/global_grid.py   → imports: numpy
planning/astar.py         → imports: numpy, heapq
planning/planner.py       → imports: planning.global_grid, planning.astar

SCRIPTS
───────
scripts/train.py
  ├── imports: models.navigation_policy.NavigationPolicy, NavigationValue
  ├── uses: gym.make() → navigation_env_cfg chain
  └── imports: skrl (PPO, SequentialTrainer)

scripts/test_reward_scales.py
  ├── uses: gym.make() → navigation_env_cfg chain
  └── dynamic task loading via importlib

scripts/cp6/visualize_generated_scenes.py
  ├── imports: scene.scene_generator (direct path, not through neurogait package)
  └── imports: scene.curriculum
  NO Isaac Sim dependency — runs with plain python3

ISOLATION NOTES
───────────────
• scene/ can be imported with plain python3 (no Isaac Sim, no pxr)
• models/ can be imported without Isaac Sim (needs torch + skrl)
• Everything in mdp/ needs Isaac Lab (isaaclab.envs, isaaclab.assets, etc.)
• Never import through neurogait.__init__ in standalone scripts
  (it cascades into isaaclab_tasks → isaaclab → pxr → GPU drivers)
""")


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

heading("Section 10 — Quick Reference", level=1)

heading("Key Environment Attributes (set on env object at runtime)", level=2)
code("""\
env._cp5_waypoints     (E, W, 2) float32  world coords of full A*/gen path
env._cp5_wp_idx        (E,)      long     current waypoint index per env
env._cp5_prev_dist     (E,)      float32  distance to wp from previous step
env._cp5_prev_action   (E, 3)    float32  previous nav action [vx,vy,hdg]
env._cp5_pos_history   (E,20,2)  float32  position ring buffer for stuck detection
env._cp5_pos_hist_idx  int       current write index into ring buffer
env._cp6_prev_action_1 (E, 3)    float32  for 2nd-order smoothness (CP6)
env._cp6_prev_action_2 (E, 3)    float32  for 2nd-order smoothness (CP6)
env._curriculum        NavigationCurriculum  scene difficulty tracker (CP6.5)
""")

heading("Key Config Parameters", level=2)
table(
    ["Parameter", "Value", "Where", "Why"],
    [
        ["decimation", "40", "navigation_env_cfg.py", "Nav policy runs at 5Hz (50Hz/40)"],
        ["episode_length_s", "120.0", "CP6EnvCfg", "6 minutes per episode"],
        ["grid_size_m", "8.0", "observations.py", "8m × 8m occupancy coverage"],
        ["n_cells", "40", "observations.py", "0.2m/cell, 1600 total"],
        ["robot_row", "10", "observations.py", "2m behind robot, 6m ahead (asymmetric)"],
        ["wp_advance", "0.3m", "observations.py", "advance waypoint when within 0.3m"],
        ["goal_threshold", "0.1m", "terminations.py", "tight goal: 0.1m from last waypoint"],
        ["ramp_iterations", "40_000", "events.py", "≈2000 training iters for full difficulty"],
    ],
    [2.0, 1.0, 2.0, 2.5],
)

heading("Training Checklist", level=2)
body("1. Run test_reward_scales.py first — confirm no term is >10× others")
body("2. Run visualize_generated_scenes.py — confirm obstacles are near path")
body("3. Start training: python3 scripts/train.py --task NeuroGait-Navigation-CP65-v0 --num_envs 512 --headless")
body("4. Watch first 50 iterations: navigation_core should climb, goal_reached near 0 initially")
body("5. The [CP6.5] log line (every 60s) shows curriculum progress")
body("6. At 2000 iterations (~40k resets): full difficulty reached")


# ─── save ────────────────────────────────────────────────────────────────────

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
